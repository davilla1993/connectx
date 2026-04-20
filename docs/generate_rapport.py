"""
Génère le rapport de projet ConnectX au format Word (.docx).
Exécuter : python docs/generate_rapport.py
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

CAPTURES = 'docs/captures'

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_font(run, name='Calibri', size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        if color:
            run.font.color.rgb = RGBColor(*color)
        run.font.name = 'Calibri'
    return p

def add_body(doc, text, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=size)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_paragraph(doc, text='', bold=False, italic=False, size=11,
                  align=None, color=None):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def add_page_break(doc):
    doc.add_page_break()

def add_image(doc, filename, caption, width_inches=5.5):
    """Insère une image centrée avec sa légende sous forme italique."""
    filepath = os.path.join(CAPTURES, filename)
    if not os.path.exists(filepath):
        # Espace réservé si la capture est manquante
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'[ CAPTURE MANQUANTE : {filename} ]')
        set_font(run, size=10, italic=True, color=(200, 0, 0))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(filepath, width=Inches(width_inches))

    # Légende
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(caption)
    set_font(run2, size=9, italic=True, color=(100, 116, 139))
    p2.paragraph_format.space_after = Pt(12)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '0F172A')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = val
            for run in cells[c_idx].paragraphs[0].runs:
                run.font.size = Pt(10)
    return table

# ── Document ───────────────────────────────────────────────────────────────────

doc = Document()
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

ACCENT  = (244, 63, 94)
PRIMARY = (15, 23, 42)
GRAY    = (100, 116, 139)

# ════════════════════════════════════════════════════════════════════════════════
# PAGE DE COUVERTURE
# ════════════════════════════════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

for text, size, bold, color in [
    ('M1 IA & BIG DATA',              13, True,  PRIMARY),
    ('Projet Python — Django',         12, False, GRAY),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ConnectX')
set_font(run, size=40, bold=True, color=ACCENT)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Réseau Social Avancé')
set_font(run, size=20, bold=True, color=PRIMARY)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Rapport de Projet')
set_font(run, size=14, italic=True, color=GRAY)

for _ in range(5):
    doc.add_paragraph()

for text, size, bold, color in [
    ('Réalisé par',                    12, False, PRIMARY),
    ('GBOSSOU Folly Sitou Carlo',      13, True,  PRIMARY),
    ('POTCHO ESSOSSOLAM Maxime',       13, True,  PRIMARY),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Encadrant : Mr. XXXXXXXX')
set_font(run, size=12, color=PRIMARY)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Avril 2026')
set_font(run, size=12, color=GRAY)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# RÉSUMÉ
# ════════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Résumé', level=1, color=ACCENT)

add_body(doc, (
    "ConnectX est une plateforme de réseau social avancé développée dans le cadre "
    "du module Python/Django du Master 1 en Intelligence Artificielle et Big Data. "
    "Le projet vise à concevoir et implémenter une application web complète reproduisant "
    "les fonctionnalités essentielles des grands réseaux sociaux modernes, tout en "
    "mettant en pratique les compétences acquises en développement web Python."
))
add_body(doc, (
    "La plateforme ConnectX permet aux utilisateurs de créer un compte, publier du contenu "
    "textuel et multimédia, suivre d'autres utilisateurs, échanger des messages en temps réel "
    "via WebSocket, publier des stories éphémères de 24 heures et recevoir des notifications "
    "instantanées pour chaque interaction sociale. L'application est construite sur Django 5, "
    "Django Channels pour la communication temps réel, et PostgreSQL comme base de données relationnelle."
))
add_body(doc, (
    "Ce rapport décrit l'ensemble de la démarche de conception et de développement : "
    "de l'analyse des besoins jusqu'aux tests de validation, en passant par les choix "
    "architecturaux et les solutions aux problèmes rencontrés."
))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Technologies principales : ')
set_font(run, size=11, bold=True)
run2 = p.add_run(
    'Python 3.12, Django 5, Django Channels, Daphne (ASGI), '
    'PostgreSQL, Bootstrap 5, WebSocket.'
)
set_font(run2, size=11)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES
# ════════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Table des matières', level=1, color=ACCENT)

toc_items = [
    ('1.', 'Introduction'),
    ('   1.1', 'Contexte et problématique'),
    ('   1.2', 'Objectifs du projet'),
    ('   1.3', 'Plan du rapport'),
    ('2.', 'Cahier des charges'),
    ('   2.1', 'Besoins fonctionnels'),
    ('   2.2', 'Besoins non fonctionnels'),
    ('   2.3', 'Contraintes techniques'),
    ('   2.4', 'Public cible'),
    ('3.', 'Conception'),
    ('   3.1', 'Architecture globale'),
    ('   3.2', 'Modèle de données'),
    ('   3.3', 'Architecture ASGI et WebSocket'),
    ('   3.4', 'Choix technologiques et justifications'),
    ('4.', 'Développement'),
    ('   4.1', 'Découpage en modules'),
    ('   4.2', 'Méthodologie de travail'),
    ('   4.3', 'Développement des fonctionnalités'),
    ('   4.4', 'Problèmes rencontrés et solutions'),
    ('5.', 'Tests et validation'),
    ('   5.1', 'Méthodologie de test'),
    ('   5.2', 'Résultats des tests'),
    ('   5.3', 'Améliorations possibles'),
    ('6.', 'Conclusion et perspectives'),
    ('Annexes', ''),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{num}  {title}')
    set_font(run, size=11, bold=(not num.startswith('   ')))

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════════

add_heading(doc, '1. Introduction', level=1, color=ACCENT)
add_heading(doc, '1.1 Contexte et problématique', level=2, color=PRIMARY)

add_body(doc, (
    "Les réseaux sociaux occupent aujourd'hui une place centrale dans la communication "
    "numérique. Qu'il s'agisse d'Instagram, de Twitter/X, de Facebook ou de LinkedIn, "
    "ces plateformes ont en commun une architecture web complexe capable de gérer de "
    "multiples interactions sociales en temps réel : publications, réactions, messagerie "
    "instantanée, notifications push. La conception d'un tel système constitue un excellent "
    "exercice d'ingénierie logicielle, car elle mobilise simultanément de nombreuses "
    "compétences : modélisation de données relationnelles complexes, programmation orientée "
    "objet, protocoles de communication réseau, gestion de l'authentification et de la "
    "sécurité, optimisation des requêtes, et conception d'interfaces utilisateur réactives."
))
add_body(doc, (
    "Dans ce contexte, le présent projet — ConnectX — répond à la problématique suivante : "
    "comment concevoir et développer, avec le framework Django, une plateforme de réseau "
    "social fonctionnelle qui gère efficacement des relations utilisateurs complexes tout "
    "en offrant des interactions en temps réel ? Cette problématique soulève plusieurs "
    "questions techniques sous-jacentes : comment modéliser proprement des relations "
    "many-to-many comme les abonnements ou les likes ? Comment implémenter la communication "
    "bidirectionnelle en temps réel avec Django, traditionnellement synchrone ? Comment "
    "garantir la sécurité des données et des échanges ?"
))
add_body(doc, (
    "Le sujet impose l'utilisation de Django Channels, une extension de Django permettant "
    "de gérer des protocoles asynchrones comme les WebSockets. Ce choix technologique "
    "impose de sortir du modèle requête-réponse classique pour adopter une architecture "
    "ASGI (Asynchronous Server Gateway Interface), ce qui constitue un vrai défi de "
    "conception pour des étudiants en Master 1."
))

add_heading(doc, '1.2 Objectifs du projet', level=2, color=PRIMARY)

add_body(doc, "Le projet ConnectX poursuit les objectifs suivants :")
add_bullet(doc, "Implémenter un système complet de publications avec support multimédia (texte + images multiples).")
add_bullet(doc, "Gérer des relations d'abonnement bidirectionnelles entre utilisateurs (follow/unfollow).")
add_bullet(doc, "Développer un système de messagerie privée fonctionnel en mode HTTP et temps réel via WebSocket.")
add_bullet(doc, "Mettre en place un système de notifications instantanées déclenché par les interactions sociales.")
add_bullet(doc, "Implémenter les stories : contenu éphémère qui expire automatiquement après 24 heures.")
add_bullet(doc, "Assurer la sécurité de l'application : authentification, permissions, protection CSRF et XSS.")
add_bullet(doc, "Concevoir une architecture Django propre, maintenable et extensible.")

add_heading(doc, '1.3 Plan du rapport', level=2, color=PRIMARY)

add_body(doc, (
    "Ce rapport est structuré en six parties principales. La section 2 présente le cahier "
    "des charges avec les besoins fonctionnels et non fonctionnels. La section 3 détaille "
    "les choix de conception et l'architecture technique. La section 4 décrit le "
    "développement par modules et la méthodologie adoptée. La section 5 présente les tests "
    "et leur validation. La section 6 conclut avec le bilan et les perspectives d'évolution. "
    "Des annexes complètent le rapport avec des extraits de code significatifs."
))

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# 2. CAHIER DES CHARGES
# ════════════════════════════════════════════════════════════════════════════════

add_heading(doc, '2. Cahier des charges', level=1, color=ACCENT)
add_heading(doc, '2.1 Besoins fonctionnels', level=2, color=PRIMARY)

add_body(doc, (
    "Les besoins fonctionnels décrivent ce que le système doit faire du point de vue de "
    "l'utilisateur. Ils sont organisés en modules correspondant aux grandes fonctionnalités "
    "de la plateforme."
))

add_paragraph(doc, 'BF1 — Gestion des comptes utilisateurs', bold=True, size=11)
add_bullet(doc, "L'utilisateur peut créer un compte avec un nom d'utilisateur unique, une adresse e-mail et un mot de passe.")
add_bullet(doc, "L'authentification se fait par e-mail et mot de passe.")
add_bullet(doc, "L'utilisateur peut modifier son profil : biographie, avatar, image de couverture, localisation et site web.")
add_bullet(doc, "L'utilisateur peut consulter le profil public d'autres utilisateurs.")
add_bullet(doc, "Un administrateur dispose d'un accès à un panneau d'administration dédié.")

add_paragraph(doc, 'BF2 — Publications', bold=True, size=11)
add_bullet(doc, "L'utilisateur peut créer une publication composée d'un texte (2 000 caractères max) et jusqu'à 5 images.")
add_bullet(doc, "Les images sont validées côté client et côté serveur (JPEG, PNG, GIF, WebP ; taille max : 5 Mo).")
add_bullet(doc, "L'utilisateur voit un fil d'actualité chronologique des publications des personnes qu'il suit.")
add_bullet(doc, "L'utilisateur peut liker ou unliker une publication (mise à jour instantanée sans rechargement).")
add_bullet(doc, "L'utilisateur peut commenter et supprimer ses propres publications.")

add_paragraph(doc, 'BF3 — Abonnements (Follow)', bold=True, size=11)
add_bullet(doc, "L'utilisateur peut suivre ou ne plus suivre d'autres utilisateurs.")
add_bullet(doc, "La liste des abonnés et des abonnements est consultable sur chaque profil.")
add_bullet(doc, "Un système de suggestions propose des personnes à suivre.")
add_bullet(doc, "Un utilisateur ne peut pas se suivre lui-même.")

add_paragraph(doc, 'BF4 — Messagerie privée', bold=True, size=11)
add_bullet(doc, "L'utilisateur peut initier une conversation privée avec n'importe quel autre utilisateur.")
add_bullet(doc, "La messagerie fonctionne en temps réel via WebSocket.")
add_bullet(doc, "Un indicateur de frappe signale quand l'interlocuteur est en train d'écrire.")
add_bullet(doc, "Les messages non lus sont comptabilisés et affichés sous forme de badge dans la navbar.")

add_paragraph(doc, 'BF5 — Notifications en temps réel', bold=True, size=11)
add_bullet(doc, "L'utilisateur reçoit une notification instantanée (toast + badge) pour chaque like, commentaire, nouveau follower ou message.")
add_bullet(doc, "Une page de liste des notifications permet de consulter l'historique complet.")
add_bullet(doc, "Les notifications sont marquées comme lues automatiquement à l'ouverture de la liste.")

add_paragraph(doc, 'BF6 — Stories', bold=True, size=11)
add_bullet(doc, "L'utilisateur peut publier une story (image + légende) visible 24 heures.")
add_bullet(doc, "Un strip horizontal dans le fil d'actualité affiche les stories actives des personnes suivies.")
add_bullet(doc, "Cliquer sur une story l'ouvre en plein écran avec une barre de progression de 10 secondes.")
add_bullet(doc, "L'auteur peut supprimer sa story avant son expiration naturelle.")

add_paragraph(doc, 'BF7 — Recherche', bold=True, size=11)
add_bullet(doc, "L'utilisateur peut rechercher d'autres utilisateurs par nom d'utilisateur ou adresse e-mail.")

add_heading(doc, '2.2 Besoins non fonctionnels', level=2, color=PRIMARY)

add_paragraph(doc, 'Sécurité', bold=True, size=11)
add_bullet(doc, "Toutes les vues sociales exigent une authentification. Un utilisateur non connecté est redirigé vers la page de connexion.")
add_bullet(doc, "Les formulaires sont protégés contre les attaques CSRF via les tokens Django.")
add_bullet(doc, "Les contenus affichés en JavaScript sont échappés pour prévenir les injections XSS.")
add_bullet(doc, "Les identifiants internes ne sont jamais exposés dans les URLs : on utilise des UUID publics.")
add_bullet(doc, "Les mots de passe sont hachés par Django (PBKDF2 avec SHA-256).")

add_paragraph(doc, 'Performance', bold=True, size=11)
add_bullet(doc, "Les requêtes sont optimisées via select_related et prefetch_related pour éviter le problème N+1.")
add_bullet(doc, "Les annotations SQL (Exists, Count) remplacent les calculs Python en mémoire.")
add_bullet(doc, "Les WebSockets évitent le polling périodique, réduisant la charge serveur.")

add_paragraph(doc, 'Maintenabilité', bold=True, size=11)
add_bullet(doc, "Application découpée en sept apps Django indépendantes, chacune responsable d'un domaine métier.")
add_bullet(doc, "Un modèle abstrait BaseModel centralise les champs d'audit communs.")
add_bullet(doc, "Le soft delete préserve l'intégrité des données et permet la restauration.")

add_paragraph(doc, 'Utilisabilité', bold=True, size=11)
add_bullet(doc, "Interface responsive grâce à Bootstrap 5 (mobile, tablette, desktop).")
add_bullet(doc, "Les interactions fréquentes (like, follow) fonctionnent sans rechargement de page (AJAX).")

add_heading(doc, '2.3 Contraintes techniques', level=2, color=PRIMARY)
add_bullet(doc, "Utilisation obligatoire du framework Django (Python).")
add_bullet(doc, "Utilisation de Django Channels pour la communication temps réel (WebSocket).")
add_bullet(doc, "Base de données relationnelle PostgreSQL.")
add_bullet(doc, "Support de l'upload d'images.")
add_bullet(doc, "Environnement de développement : Windows 11, Python 3.12.")

add_heading(doc, '2.4 Public cible', level=2, color=PRIMARY)
add_body(doc, (
    "ConnectX s'adresse à un public jeune et technophile, habitué aux réseaux sociaux. "
    "Dans un cadre académique, les utilisateurs sont des étudiants et enseignants qui "
    "peuvent se connecter, partager du contenu et interagir. La plateforme est conçue "
    "pour être intuitive, avec une interface moderne inspirée des grands réseaux sociaux, "
    "de façon à minimiser la courbe d'apprentissage pour tout nouvel utilisateur."
))

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# 3. CONCEPTION
# ════════════════════════════════════════════════════════════════════════════════

add_heading(doc, '3. Conception', level=1, color=ACCENT)
add_heading(doc, '3.1 Architecture globale', level=2, color=PRIMARY)

add_body(doc, (
    "ConnectX adopte une architecture modulaire fondée sur la séparation des responsabilités. "
    "L'application est découpée en sept modules Django (apps), chacun encapsulant un domaine "
    "métier précis. Cette organisation facilite la maintenance, les tests indépendants "
    "et l'évolutivité du projet."
))

add_table(doc,
    ['App Django', 'Responsabilité', 'Modèles principaux'],
    [
        ['core',          'Socle commun : modèle abstrait d\'audit',       'BaseModel'],
        ['accounts',      'Comptes et profils utilisateurs',               'User, Profile'],
        ['posts',         'Publications, images, likes, commentaires',     'Post, PostImage, Like, Comment'],
        ['social',        'Relations d\'abonnement',                       'Follow'],
        ['chat',          'Messagerie privée (HTTP + WebSocket)',           'Conversation, Message'],
        ['notifications', 'Notifications temps réel (WebSocket)',          'Notification'],
        ['stories',       'Stories éphémères (24h)',                       'Story'],
    ]
)

doc.add_paragraph()
add_body(doc, (
    "Le cœur de l'architecture repose sur deux couches de communication complémentaires : "
    "la couche HTTP classique Django (vues basées sur des classes, templates, ORM) et la "
    "couche WebSocket gérée par Django Channels. Ces deux couches coexistent grâce au "
    "protocole ASGI qui remplace WSGI lorsque des connexions persistantes sont nécessaires."
))

add_heading(doc, '3.2 Modèle de données', level=2, color=PRIMARY)

add_body(doc, (
    "Le modèle de données est le cœur du système. Toutes les entités métier héritent du "
    "modèle abstrait BaseModel qui fournit des champs d'audit communs."
))

add_paragraph(doc, 'Le modèle BaseModel', bold=True, size=11)
add_body(doc, (
    "BaseModel est une classe abstraite Django (Meta: abstract = True) dont héritent tous "
    "les modèles de ConnectX, à l'exception du modèle User. Ce pattern est inspiré du "
    "BaseEntity utilisé dans les architectures Java/Spring Data. Il apporte : "
    "public_id (UUID), created_at, updated_at, created_by, updated_by, is_deleted, "
    "deleted_at, deleted_by."
))
add_body(doc, (
    "La séparation entre l'identifiant interne (id séquentiel, jamais exposé) et le "
    "public_id (UUID) est une mesure de sécurité importante : elle empêche l'énumération "
    "d'objets par un attaquant tentant de deviner des URLs séquentielles."
))
add_body(doc, (
    "Le soft delete permet de « supprimer » un enregistrement sans l'effacer physiquement. "
    "is_deleted passe à True, et toutes les requêtes filtrent sur is_deleted=False. "
    "Cette approche préserve l'intégrité référentielle, permet la restauration en cas "
    "d'erreur et maintient un historique complet pour l'audit."
))

add_paragraph(doc, 'User et Profile', bold=True, size=11)
add_body(doc, (
    "Le modèle User étend AbstractUser en ajoutant email (unique, utilisé comme identifiant "
    "de connexion), public_id (UUID) et is_online (statut temps réel mis à jour par les "
    "WebSockets). Le Profile est lié à User par OneToOne et stocke les informations "
    "sociales : bio, avatar, image de couverture, localisation, site web. Il est créé "
    "automatiquement via signal post_save à chaque création d'utilisateur."
))

add_paragraph(doc, 'Publications (Post, PostImage, Like, Comment)', bold=True, size=11)
add_body(doc, (
    "La séparation des images dans PostImage respecte la première forme normale (1NF) "
    "et permet d'associer des métadonnées à chaque image. La relation Like utilise "
    "unique_together sur (user, post) pour garantir l'unicité en base de données. "
    "Le soft delete sur Like permet de conserver l'historique tout en autorisant le re-like."
))

add_paragraph(doc, 'Follow', bold=True, size=11)
add_body(doc, (
    "Le modèle Follow représente la relation d'abonnement avec deux clés étrangères vers "
    "User : follower et following. La contrainte unique_together empêche les doublons. "
    "Les related_name 'following_relations' et 'follower_relations' permettent de naviguer "
    "efficacement dans les deux sens depuis un objet User."
))

add_paragraph(doc, 'Conversation et Message', bold=True, size=11)
add_body(doc, (
    "Conversation est lié aux participants via ManyToMany, supportant techniquement les "
    "groupes de plus de deux personnes. Message référence l'expéditeur et la conversation. "
    "Le champ is_read permet de comptabiliser les messages non lus via un context processor "
    "injecté dans tous les templates."
))

add_paragraph(doc, 'Notification', bold=True, size=11)
add_body(doc, (
    "Le modèle Notification centralise les quatre types d'événements sociaux : like, "
    "commentaire, abonnement, message. Il référence le destinataire, l'expéditeur, "
    "le type d'événement, une publication optionnelle et l'état de lecture. "
    "La méthode get_message() retourne le texte localisé affiché à l'utilisateur."
))

add_paragraph(doc, 'Story', bold=True, size=11)
add_body(doc, (
    "Le modèle Story stocke une image, une légende et une date d'expiration calculée "
    "automatiquement (now + 24h) en surchargeant save(). Les stories sont filtrées "
    "directement dans la requête ORM via expires_at__gt=timezone.now(), ce qui est "
    "plus efficace qu'un filtrage Python post-récupération."
))

add_heading(doc, '3.3 Architecture ASGI et WebSocket', level=2, color=PRIMARY)

add_body(doc, (
    "Le fichier asgi.py configure un ProtocolTypeRouter qui aiguille les connexions selon "
    "leur protocole : les requêtes HTTP sont traitées par Django classique, tandis que "
    "les connexions WebSocket sont acheminées vers un URLRouter Channels. "
    "L'AuthMiddlewareStack injecte automatiquement l'objet User dans le scope WebSocket."
))
add_body(doc, "ConnectX expose deux endpoints WebSocket distincts :")
add_bullet(doc, "ws/chat/<conversation_id>/ — ChatConsumer pour la messagerie temps réel.")
add_bullet(doc, "ws/notifications/ — NotificationConsumer pour les notifications push.")
add_body(doc, (
    "Les signaux Django sont synchrones, mais l'envoi au channel layer est asynchrone. "
    "La fonction async_to_sync de asgiref permet de « ponter » ces deux mondes : elle "
    "exécute une coroutine depuis un contexte synchrone, ce qui est exactement ce dont "
    "on a besoin dans un signal. Le channel layer InMemoryChannelLayer est utilisé en "
    "développement. En production, il faudra migrer vers channels_redis avec Redis."
))

add_heading(doc, '3.4 Choix technologiques et justifications', level=2, color=PRIMARY)

for titre, texte in [
    ('Python et Django 5',
     "Python s'impose naturellement dans le contexte d'un Master en IA et Big Data. Django "
     "est le framework web Python le plus mature, suivant le principe « batteries included » : "
     "ORM, authentification, migrations, formulaires, template engine — tout est intégré. "
     "Django 5 apporte des améliorations de performance et une meilleure intégration ASGI native."),
    ('Django Channels et Daphne',
     "Django Channels est la solution officielle pour ajouter le support WebSocket à Django. "
     "Elle s'intègre nativement avec l'écosystème Django et suit les mêmes conventions, "
     "réduisant la courbe d'apprentissage. Daphne est le serveur ASGI de référence pour "
     "Channels : il gère la coexistence des connexions HTTP et WebSocket sur le même port."),
    ('PostgreSQL',
     "PostgreSQL est la base de données relationnelle open-source la plus avancée, avec "
     "un excellent support des UUID natifs et des contraintes d'intégrité complexes. "
     "Django possède un backend PostgreSQL très optimisé. C'est le standard de fait "
     "pour les applications Django en production."),
    ('Bootstrap 5 et Bootstrap Icons',
     "Bootstrap 5 permet de créer rapidement des interfaces responsive et cohérentes "
     "sans écrire beaucoup de CSS personnalisé. Bootstrap Icons complète l'interface "
     "avec une bibliothèque d'icônes SVG légère. Un CSS custom définit la charte graphique "
     "de ConnectX (palette de couleurs, rayons, ombres portées)."),
    ('Jazzmin (admin)',
     "Jazzmin est un thème d'administration Django moderne qui offre une navigation latérale, "
     "des icônes par modèle, un mode sombre et une configuration déclarative. Ce choix "
     "améliore significativement l'expérience des administrateurs de la plateforme."),
]:
    add_paragraph(doc, titre, bold=True, size=11)
    add_body(doc, texte)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# 4. DÉVELOPPEMENT
# ════════════════════════════════════════════════════════════════════════════════

add_heading(doc, '4. Développement', level=1, color=ACCENT)
add_heading(doc, '4.1 Découpage en modules', level=2, color=PRIMARY)

add_body(doc, (
    "Le projet est organisé en deux répertoires principaux : apps/ qui contient l'ensemble "
    "des modules Django, et templates/ qui centralise tous les templates HTML. Chaque app "
    "suit la même structure interne : models.py, views.py, urls.py, admin.py, forms.py "
    "et migrations/. Les apps chat et notifications ajoutent consumers.py et routing.py "
    "pour Django Channels."
))

add_heading(doc, '4.2 Méthodologie de travail', level=2, color=PRIMARY)

add_body(doc, (
    "Le développement de ConnectX a suivi une approche itérative et incrémentale, proche "
    "du cycle Agile. Chaque itération aboutissait à une version testable de l'application."
))
add_bullet(doc, "Itération 1 : Mise en place Django, modèle User personnalisé, authentification.")
add_bullet(doc, "Itération 2 : App posts — modèles, vues CRUD, template de feed.")
add_bullet(doc, "Itération 3 : App social — Follow, suggestions, profils.")
add_bullet(doc, "Itération 4 : App chat — messagerie HTTP puis WebSocket temps réel.")
add_bullet(doc, "Itération 5 : Extensions — notifications WebSocket et stories éphémères.")
add_bullet(doc, "Itération 6 : Polissage — admin Jazzmin, optimisation requêtes, sécurité, CSS.")
add_body(doc, (
    "Cette approche a permis de tester chaque fonctionnalité indépendamment et d'identifier "
    "les problèmes tôt. Le contrôle de version Git a été utilisé tout au long du projet "
    "pour suivre l'historique des modifications."
))

add_heading(doc, '4.3 Développement des fonctionnalités', level=2, color=PRIMARY)

# ── Auth ──────────────────────────────────────────────────────────────────────
add_paragraph(doc, 'Authentification et gestion des profils', bold=True, size=11)

add_body(doc, (
    "Le modèle User personnalisé hérite de AbstractUser et définit EMAIL comme USERNAME_FIELD. "
    "Ce choix implique d'adapter la configuration AUTH_USER_MODEL avant la première migration, "
    "car Django ne supporte pas le changement du modèle utilisateur sur une base existante. "
    "La vue CustomLoginView surcharge get_success_url() pour rediriger les administrateurs "
    "vers /admin/ et les utilisateurs normaux vers le feed. Le mixin SocialLoginRequired "
    "étend LoginRequiredMixin en ajoutant cette redirection automatique des comptes staff."
))
add_body(doc, (
    "Le profil est créé automatiquement via un signal post_save sur User, garantissant "
    "que chaque utilisateur possède toujours un profil associé, même s'il est créé via "
    "l'admin ou un script de gestion Django."
))

doc.add_paragraph()
add_image(doc, '2-inscription.png',
          "Figure 1 — Formulaire d'inscription : création de compte avec email, "
          "nom d'utilisateur et mot de passe",
          width_inches=5.5)
add_image(doc, '1-connexion.png',
          "Figure 2 — Page de connexion : authentification par e-mail et mot de passe",
          width_inches=3.8)

# ── Profil ────────────────────────────────────────────────────────────────────
add_paragraph(doc, 'Profil utilisateur', bold=True, size=11)
add_body(doc, (
    "La page de profil affiche la photo de couverture, l'avatar, la biographie, "
    "la localisation, le site web et les statistiques sociales (publications, abonnés, "
    "abonnements). L'édition du profil utilise un formulaire ModelForm lié à l'instance "
    "Profile de l'utilisateur connecté, avec upload d'avatar et d'image de couverture."
))

add_image(doc, '5-profil utilisateur.png',
          "Figure 3 — Page de profil : cover, avatar, bio et statistiques sociales",
          width_inches=6.0)

# ── Abonnements ──────────────────────────────────────────────────────────────
add_paragraph(doc, 'Abonnements et découverte', bold=True, size=11)
add_body(doc, (
    "Le système de follow repose sur un endpoint AJAX (FollowToggleView) qui bascule "
    "l'état de la relation sans rechargement de page. La vue vérifie l'existence d'un "
    "Follow entre les deux utilisateurs, puis le réactive (re-follow après unfollow), "
    "le crée (premier follow) ou le supprime logiquement (unfollow). La page Suggestions "
    "liste tous les utilisateurs que l'utilisateur connecté ne suit pas encore, en "
    "excluant les comptes staff et superuser. Dans le feed, les cinq premières suggestions "
    "sont affichées dans une sidebar avec un bouton « Suivre » AJAX."
))

add_image(doc, "10-suggestion d'amis.png",
          "Figure 4 — Page de suggestions : utilisateurs à suivre avec bouton d'abonnement rapide",
          width_inches=6.0)

# ── Feed ─────────────────────────────────────────────────────────────────────
add_paragraph(doc, 'Fil d\'actualité et publications', bold=True, size=11)
add_body(doc, (
    "Le fil d'actualité (FeedView) agrège les publications des utilisateurs suivis et les "
    "siennes dans l'ordre chronologique inverse. La requête ORM utilise select_related "
    "pour l'auteur et son profil (évite les requêtes N+1), prefetch_related pour les "
    "images, likes et commentaires, et une sous-requête annotée (Exists + OuterRef) pour "
    "déterminer si l'utilisateur a liké chaque post sans effectuer de requête supplémentaire."
))
add_body(doc, (
    "En haut du fil se trouve le strip stories : une rangée horizontale scrollable "
    "affichant les stories actives (non expirées) des personnes suivies. Si l'utilisateur "
    "n'a pas de story active, un bouton « Ajouter » est affiché à sa place."
))

add_image(doc, "3-Fil d'actualité.png",
          "Figure 5 — Fil d'actualité : strip stories en haut, publications et sidebar suggestions",
          width_inches=6.0)

add_body(doc, (
    "La validation des images est effectuée à deux niveaux : côté client en JavaScript "
    "(prévisualisation, comptage, vérification du format avant envoi) et côté serveur en "
    "Python (taille max 5 Mo, types MIME autorisés, nombre max 5). Cette double validation "
    "garantit une bonne expérience utilisateur tout en assurant la sécurité côté serveur."
))

add_image(doc, '4-Formulaire de publication.png',
          "Figure 6 — Formulaire de publication avec prévisualisation des images sélectionnées",
          width_inches=6.0)

# ── Chat ─────────────────────────────────────────────────────────────────────
add_paragraph(doc, 'Messagerie privée et WebSocket', bold=True, size=11)
add_body(doc, (
    "La messagerie combine deux couches complémentaires. La couche HTTP gère la liste "
    "des conversations, l'historique et la création d'une nouvelle conversation. "
    "La couche WebSocket (ChatConsumer) gère l'envoi et la réception en temps réel."
))
add_body(doc, (
    "Lorsqu'un utilisateur ouvre une conversation, le JavaScript établit une connexion "
    "WebSocket vers ws/chat/<id>/. Le consumer vérifie l'authentification et la "
    "participation à la conversation avant d'accepter la connexion. Les messages sont "
    "sauvegardés en base de données via database_sync_to_async, puis diffusés au groupe "
    "Channels. Les indicateurs de frappe et de présence sont également gérés via ce mécanisme."
))

add_image(doc, '6-conversation messagerie.png',
          "Figure 7 — Messagerie en temps réel : bulles de dialogue, indicateur de frappe "
          "et statut en ligne",
          width_inches=6.0)

# ── Notifications ─────────────────────────────────────────────────────────────
add_paragraph(doc, 'Notifications en temps réel', bold=True, size=11)
add_body(doc, (
    "Le système de notifications est déclenché par des signaux Django (post_save) sur "
    "Like, Comment, Follow et Message. Chaque signal crée un enregistrement Notification "
    "en base de données, puis envoie un événement au channel layer via async_to_sync. "
    "Le groupe ciblé est notif_<recipient_id>, propre à chaque utilisateur."
))
add_body(doc, (
    "Côté client, un script présent sur toutes les pages (base.html) maintient une "
    "connexion WebSocket vers ws/notifications/. À chaque notification reçue, le badge "
    "de la cloche est incrémenté et un toast discret s'affiche en bas à droite de "
    "l'écran avant de disparaître automatiquement après 4 secondes."
))

add_image(doc, '7-toast de notification.png',
          "Figure 8 — Toast de notification en temps réel : message instantané sans rechargement de page",
          width_inches=3.0)

# ── Stories ───────────────────────────────────────────────────────────────────
add_paragraph(doc, 'Stories éphémères', bold=True, size=11)
add_body(doc, (
    "Les stories disparaissent automatiquement après 24 heures sans tâche planifiée. "
    "L'expiration est gérée par filtrage ORM (expires_at__gt=timezone.now()), excluant "
    "naturellement les stories expirées. La visionneuse est une page plein écran avec "
    "l'image en fond sombre, l'auteur en en-tête, la légende en bas et une barre de "
    "progression animée sur 10 secondes, après quoi l'utilisateur est automatiquement "
    "redirigé vers le feed."
))

add_image(doc, '8-Visionneuse_de_story.png',
          "Figure 9 — Visionneuse de story plein écran : barre de progression, "
          "auteur en en-tête et légende en bas",
          width_inches=6.0)

# ── Admin ─────────────────────────────────────────────────────────────────────
add_paragraph(doc, 'Interface d\'administration Jazzmin', bold=True, size=11)
add_body(doc, (
    "L'interface d'administration utilise Jazzmin, un thème moderne qui remplace "
    "l'admin Django par défaut. Il offre une sidebar avec icônes par modèle, une "
    "navigation claire et une configuration entièrement déclarative. Tous les modèles "
    "sont enregistrés avec des vues list personnalisées (colonnes, filtres, recherche) "
    "et des champs en lecture seule pour les données d'audit. Les administrateurs "
    "sont automatiquement redirigés vers /admin/ dès leur connexion."
))

add_image(doc, '9-interface admin.png',
          "Figure 10 — Interface d'administration Jazzmin : dashboard avec tous les modèles ConnectX",
          width_inches=6.0)

add_heading(doc, '4.4 Problèmes rencontrés et solutions apportées', level=2, color=PRIMARY)

problemes = [
    (
        'Problème 1 : Encodage UTF-8 sur Windows avec PostgreSQL',
        "Sur Windows avec une locale française, psycopg2 retournait des UnicodeDecodeError "
        "lors du traitement de messages d'erreur PostgreSQL contenant des caractères accentués. "
        "Ce problème est lié à la locale système Windows (CP1252) alors que PostgreSQL utilise UTF-8.",
        "Forcer les variables d'environnement PGCLIENTENCODING=UTF8 et PGSSLMODE=prefer "
        "au niveau des fichiers asgi.py et settings.py, avant toute importation Django. "
        "Solution simple et non intrusive, sans modifier la configuration système."
    ),
    (
        'Problème 2 : Signal synchrone ↔ channel layer asynchrone',
        "Les signaux Django sont synchrones et s'exécutent dans le même thread que la "
        "requête HTTP. Or, l'envoi d'un message au channel layer est une opération "
        "asynchrone. Appeler directement une coroutine depuis un signal est impossible "
        "sans event loop actif.",
        "Utiliser async_to_sync de asgiref, qui crée temporairement un event loop pour "
        "exécuter la coroutine de manière synchrone. C'est la méthode recommandée "
        "officiellement dans la documentation Django Channels."
    ),
    (
        'Problème 3 : N+1 queries dans le feed',
        "Une première version du feed chargeait les posts sans optimisation, entraînant "
        "une requête SQL par post pour récupérer l'avatar de l'auteur — soit N+1 requêtes "
        "pour N posts. Sur un feed avec 50 publications, cela représentait plus de 50 "
        "aller-retours vers la base de données pour une seule page.",
        "Utiliser select_related('author', 'author__profile') pour charger l'auteur et "
        "son profil en une seule requête SQL avec JOIN, et prefetch_related pour les "
        "relations inversées. L'annotation user_liked via Exists + OuterRef évite de "
        "charger tous les likes en mémoire."
    ),
    (
        'Problème 4 : Gestion des fichiers multiples dans le formulaire HTML',
        "Si l'utilisateur sélectionne des images en plusieurs fois, le champ file input "
        "ne conserve que la dernière sélection, écrasant les précédentes.",
        "Utiliser un objet DataTransfer JavaScript pour maintenir la liste des fichiers "
        "sélectionnés. À chaque nouvelle sélection, les fichiers sont ajoutés à la liste "
        "interne (limitée à 5) et un DataTransfer synthétique est reconstruit et assigné "
        "à fileInput.files."
    ),
    (
        'Problème 5 : Sécurité des consumers WebSocket',
        "Sans précaution, un WebSocket peut être connecté par n'importe qui, même un "
        "utilisateur non authentifié. Il était crucial de valider l'identité à la connexion.",
        "Dans connect() de chaque consumer, vérifier que scope['user'].is_authenticated. "
        "Pour ChatConsumer, vérifier en plus que l'utilisateur est participant de la "
        "conversation. En cas d'échec, appeler self.close() immédiatement."
    ),
]

for titre, pb, sol in problemes:
    add_paragraph(doc, titre, bold=True, size=11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('Problème : ')
    set_font(run, size=11, italic=True)
    run2 = p.add_run(pb)
    set_font(run2, size=11)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run3 = p2.add_run('Solution : ')
    set_font(run3, size=11, bold=True, color=ACCENT)
    run4 = p2.add_run(sol)
    set_font(run4, size=11)
    doc.add_paragraph()

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# 5. TESTS ET VALIDATION
# ════════════════════════════════════════════════════════════════════════════════

add_heading(doc, '5. Tests et validation', level=1, color=ACCENT)
add_heading(doc, '5.1 Méthodologie de test adoptée', level=2, color=PRIMARY)

add_body(doc, (
    "Les tests de ConnectX ont été réalisés principalement par des tests fonctionnels "
    "manuels et des tests de sécurité, complétés par la commande de vérification "
    "intégrée de Django."
))

add_paragraph(doc, 'Tests de vérification système Django', bold=True, size=11)
add_body(doc, (
    "La commande python manage.py check analyse statiquement l'ensemble de la configuration : "
    "modèles, migrations, paramètres, conflits de related_name. Elle a été exécutée "
    "systématiquement après chaque ajout de modèle. Elle a notamment permis de détecter "
    "des conflits de related_name dans BaseModel, résolu grâce à la syntaxe %(class)s."
))

add_paragraph(doc, 'Tests fonctionnels manuels', bold=True, size=11)
add_body(doc, (
    "Chaque fonctionnalité a été testée en suivant des scénarios utilisateur complets "
    "avec plusieurs comptes de test, couvrant le chemin nominal (happy path) et les cas "
    "d'erreur principaux (formulaire invalide, permissions manquantes, ressource inexistante)."
))

add_paragraph(doc, 'Tests de sécurité', bold=True, size=11)
add_body(doc, (
    "Une série de tests basiques ont été effectués : accès aux pages sociales sans "
    "authentification, tentative de suppression d'une publication appartenant à un "
    "autre utilisateur, connexion à un WebSocket de conversation sans y être participant."
))

add_heading(doc, '5.2 Résultats des tests', level=2, color=PRIMARY)

add_table(doc,
    ['Fonctionnalité', 'Scénario testé', 'Résultat'],
    [
        ['Inscription',       'Compte avec email valide',              'Réussi'],
        ['Inscription',       'Email déjà utilisé',                    'Erreur affichée'],
        ['Connexion',         'Email + mdp corrects',                  'Réussi — feed'],
        ['Connexion',         'Mot de passe incorrect',                'Erreur Django standard'],
        ['Publication',       'Texte + 3 images valides',              'Réussi — grille images'],
        ['Publication',       '6 images (dépasse max)',                'Erreur client bloquée'],
        ['Publication',       'Image > 5 Mo',                         'Erreur serveur retournée'],
        ['Like',              'Like puis unlike',                      'Compteur AJAX OK'],
        ['Commentaire',       'Texte valide',                          'Affiché sur post_detail'],
        ['Follow',            'Suivre un utilisateur',                 'Suggestion disparaît'],
        ['Follow',            'Se suivre soi-même',                   'Bloqué par la vue'],
        ['Chat HTTP',         'Démarrer conversation',                 'Créée ou retrouvée'],
        ['Chat WebSocket',    'Envoi/réception temps réel',            'Réussi — 2 onglets'],
        ['Chat WebSocket',    'Indicateur de frappe',                  'Réussi'],
        ['Chat WebSocket',    'Connexion sans authentification',       'Refusée (close())'],
        ['Chat WebSocket',    'Connexion hors conversation',           'Refusée'],
        ['Notifications',     'Like → toast destinataire',            'Reçu temps réel'],
        ['Notifications',     'Follow → notification',                 'Réussi'],
        ['Notifications',     'Liste — marquage lu auto',              'Réussi'],
        ['Stories',           'Création avec image',                   'Dans le strip feed'],
        ['Stories',           'Visionneuse plein écran',               'Barre progression OK'],
        ['Stories',           'Expiration automatique 24h',            'Disparaît du feed'],
        ['Stories',           'Suppression par l\'auteur',             'Réussi'],
        ['Sécurité',          'Accès feed sans login',                 'Redirection /login/'],
        ['Sécurité',          'DELETE post d\'un autre',               '404 get_object_or_404'],
        ['Admin',             'Connexion compte staff',                'Redirection /admin/'],
    ]
)

add_heading(doc, '5.3 Améliorations possibles', level=2, color=PRIMARY)

for titre, texte in [
    ('Tests automatisés',
     "L'absence de tests unitaires et d'intégration automatisés est la principale lacune. "
     "Il serait pertinent d'ajouter des TestCase Django pour chaque vue et modèle, ainsi "
     "que des tests de consumer WebSocket via ChannelsLiveServerTestCase."),
    ('Pagination',
     "Le feed et la liste des notifications ne sont pas paginés. Sur une base volumineuse, "
     "charger plusieurs centaines de publications en une requête deviendrait problématique. "
     "Une pagination classique ou infinie est une amélioration prioritaire."),
    ('Channel layer Redis en production',
     "InMemoryChannelLayer ne supporte pas plusieurs processus. En production avec un "
     "load balancer, les messages WebSocket ne pourraient pas être partagés entre workers. "
     "La migration vers channels_redis est indispensable."),
    ('Nettoyage des stories expirées',
     "Les stories ne sont pas supprimées physiquement à expiration. Une tâche planifiée "
     "(Celery Beat ou django-crontab) pourrait effectuer un nettoyage périodique."),
    ('Conversations de groupe',
     "Le modèle Conversation supporte techniquement les groupes via ManyToMany. "
     "Les vues et l'interface ne le supportent pas encore, mais l'extension est préparée."),
]:
    add_paragraph(doc, titre, bold=True, size=11)
    add_body(doc, texte)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# 6. CONCLUSION ET PERSPECTIVES
# ════════════════════════════════════════════════════════════════════════════════

add_heading(doc, '6. Conclusion et perspectives', level=1, color=ACCENT)
add_heading(doc, '6.1 Bilan global du projet', level=2, color=PRIMARY)

add_body(doc, (
    "Le projet ConnectX a atteint l'ensemble de ses objectifs. La plateforme implémente "
    "les six fonctionnalités principales du sujet — publications avec images, abonnements, "
    "messagerie, Django Channels, upload d'images et relations complexes — ainsi que les "
    "deux extensions demandées : notifications en temps réel et stories éphémères."
))
add_body(doc, (
    "Sur le plan technique, ce projet a permis de consolider la maîtrise de Django ORM, "
    "des vues basées sur des classes et du système de templates. L'intégration de Django "
    "Channels a représenté le défi le plus significatif : sortir du modèle synchrone "
    "de Django pour adopter une architecture ASGI asynchrone nécessite de comprendre "
    "des concepts nouveaux — consumers, channel groups, concurrence async/await."
))
add_body(doc, (
    "La conception du modèle de données a également été enrichissante : le pattern "
    "BaseModel avec soft delete et audit trail, la séparation entre identifiant interne "
    "et UUID public, et l'utilisation de signaux Django pour découpler les effets de bord "
    "sont des pratiques professionnelles directement transposables en entreprise."
))
add_body(doc, (
    "Sur le plan méthodologique, l'approche itérative et incrémentale s'est révélée "
    "efficace. Elle a permis de livrer régulièrement des versions fonctionnelles et "
    "d'ajuster le plan en fonction des difficultés rencontrées, plutôt que de les "
    "découvrir en fin de projet."
))

add_heading(doc, '6.2 Pistes d\'amélioration et évolutions possibles', level=2, color=PRIMARY)

add_body(doc, (
    "ConnectX, dans sa version actuelle, est une démonstration fonctionnelle solide. "
    "Pour évoluer vers une plateforme prête pour la production, plusieurs évolutions "
    "seraient nécessaires :"
))
add_bullet(doc, "Déploiement : containerisation Docker, Nginx comme reverse proxy, Daphne en mode production, Redis pour le channel layer, PostgreSQL managé.")
add_bullet(doc, "Scalabilité : pagination du feed, cache Redis pour les requêtes fréquentes, CDN pour les fichiers media.")
add_bullet(doc, "Fonctionnalités : partage de publications, mentions (@username), hashtags, réactions variées, appels vidéo via WebRTC.")
add_bullet(doc, "Modération : signalement de contenu, système de blocage d'utilisateurs, file d'attente de modération.")
add_bullet(doc, "API REST ou GraphQL (Django REST Framework) pour connecter une application mobile React Native ou Flutter.")
add_bullet(doc, "Accessibilité : conformité WCAG 2.1 (textes alternatifs, navigation clavier, contrastes).")

add_body(doc, (
    "Ce projet a été une expérience de développement complète et formatrice, alliant "
    "conception technique rigoureuse, résolution de problèmes concrets et apprentissage "
    "de technologies avancées. ConnectX constitue une base solide sur laquelle des "
    "fonctionnalités additionnelles pourraient être construites dans le cadre d'un "
    "projet de Master 2 ou d'un projet personnel."
))

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# ANNEXES
# ════════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Annexes', level=1, color=ACCENT)

annexes = [
    ('Annexe A — BaseModel (core/models.py)',
     'Le modèle abstrait dont héritent toutes les entités ConnectX :',
     """\
class BaseModel(models.Model):
    public_id  = models.UUIDField(default=uuid.uuid4, editable=False,
                                   unique=True, db_index=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    created_by = models.ForeignKey(AUTH_USER_MODEL, null=True, blank=True,
                     on_delete=SET_NULL, related_name='%(class)s_created')
    updated_by = models.ForeignKey(AUTH_USER_MODEL, null=True, blank=True,
                     on_delete=SET_NULL, related_name='%(class)s_updated')
    is_deleted = models.BooleanField(default=False)
    deleted_at = models.DateTimeField(null=True, blank=True)
    deleted_by = models.ForeignKey(AUTH_USER_MODEL, null=True, blank=True,
                     on_delete=SET_NULL, related_name='%(class)s_deleted')
    class Meta:
        abstract = True"""),

    ('Annexe B — ChatConsumer (chat/consumers.py)',
     'Le consumer WebSocket gérant la messagerie en temps réel :',
     """\
class ChatConsumer(AsyncWebsocketConsumer):
    async def connect(self):
        self.conversation_id = self.scope['url_route']['kwargs']['conversation_id']
        self.room_group_name = f'chat_{self.conversation_id}'
        self.user = self.scope['user']
        if not self.user.is_authenticated:
            await self.close(); return
        if not await self.is_participant():
            await self.close(); return
        await self.channel_layer.group_add(self.room_group_name, self.channel_name)
        await self.accept()

    async def receive(self, text_data):
        data = json.loads(text_data)
        content = data.get('message', '').strip()
        if not content:
            return
        message = await self.save_message(content)
        await self.channel_layer.group_send(self.room_group_name, {
            'type': 'chat_message',
            'message': content,
            'sender': self.user.username,
            'created_at': message.created_at.isoformat(),
        })

    @database_sync_to_async
    def save_message(self, content):
        conversation = Conversation.objects.get(id=self.conversation_id)
        return Message.objects.create(
            conversation=conversation, sender=self.user,
            content=content, created_by=self.user
        )"""),

    ('Annexe C — Signaux de notifications (notifications/signals.py)',
     'Déclenchement automatique des notifications par signaux Django :',
     """\
def push_notification(notification):
    channel_layer = get_channel_layer()
    async_to_sync(channel_layer.group_send)(
        f'notif_{notification.recipient_id}',
        {'type': 'new_notification',
         'id': notification.id,
         'notif_type': notification.notif_type,
         'message': notification.get_message()}
    )

@receiver(post_save, sender='posts.Like')
def on_like(sender, instance, created, **kwargs):
    if not created or instance.user == instance.post.author:
        return
    notif = Notification.objects.create(
        recipient=instance.post.author, sender=instance.user,
        notif_type=Notification.TYPE_LIKE, post=instance.post,
        created_by=instance.user,
    )
    push_notification(notif)"""),

    ('Annexe D — Modèle Story avec expiration automatique',
     'Logique d\'expiration automatique à 24 heures :',
     """\
class Story(BaseModel):
    author     = models.ForeignKey(User, on_delete=CASCADE, related_name='stories')
    image      = models.ImageField(upload_to='stories/')
    caption    = models.CharField(max_length=200, blank=True)
    expires_at = models.DateTimeField()

    def save(self, *args, **kwargs):
        if not self.pk:
            self.expires_at = timezone.now() + timedelta(hours=24)
        super().save(*args, **kwargs)

    @property
    def is_expired(self):
        return timezone.now() > self.expires_at"""),

    ('Annexe E — Configuration ASGI (connectx/asgi.py)',
     'Point d\'entrée ASGI unifiant connexions HTTP et WebSocket :',
     """\
application = ProtocolTypeRouter({
    'http': get_asgi_application(),
    'websocket': AuthMiddlewareStack(
        URLRouter(chat_ws + notif_ws)
    ),
})"""),

    ('Annexe F — Optimisation requête FeedView',
     'Requête ORM optimisée évitant le problème N+1 :',
     """\
liked_subquery = Like.objects.filter(
    post=OuterRef('pk'), user=request.user, is_deleted=False
)
posts = (
    Post.objects
    .filter(is_deleted=False,
            author_id__in=following_ids + [request.user.id])
    .select_related('author', 'author__profile')
    .prefetch_related('images', 'likes', 'comments')
    .annotate(user_liked=Exists(liked_subquery))
    .order_by('-created_at')
)"""),
]

for titre, intro, code in annexes:
    add_heading(doc, titre, level=2, color=PRIMARY)
    add_body(doc, intro)
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    doc.add_paragraph()

# ── Sauvegarde ────────────────────────────────────────────────────────────────

output = 'docs/Rapport_ConnectX.docx'
doc.save(output)
print(f'Rapport généré : {output}')
